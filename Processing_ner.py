import  pandas  as pd
from docx import Document
import re
import os
import jieba

def build_csv(excel_path, word_path, save_path):  
    """建立dataframe
    path_csv: excel文件夹路径
    save_path: 需要存储的文件夹路径

    """
    document = None
    dfall = None
    nums = len(os.listdir(excel_path))
    print(nums)
    num = 0
    for info in os.listdir(excel_path):
        if info.endswith(".xlsx"):             
            domain = os.path.abspath(excel_path) #获取文件夹的路径
            info = os.path.join(domain,info) #将路径与文件名结合起来就是每个文件的完整路径 
            df = pd.read_excel(info) 
            dfall= pd.concat([dfall,df],axis=0,ignore_index=True,sort=False) #拼接dataframe
            num+=1
            print("\r%.2f%%" % ((num / nums)*100), end='')

    dfall = dfall.drop(columns=['基本案情', '争议焦点', '法院认为', '裁判结果', '链接'])
    dfall= dfall.dropna(axis=0) #去掉有空字段的数据
    dfall = dfall.drop_duplicates(subset=['案例名称'])#去重

    # for index,i in list(enumerate(dfall.loc[:,"审判日期"])):#改变时间格式
    #     timelist = dfall.loc[index,"审判日期"].split('-')
    #     dfall.loc[index,"审判日期"] = time2chinese(timelist)
    dfall=pd.concat([dfall, pd.DataFrame(columns=["原告","被告"])],sort=False)
    counts = len(dfall)
    count = 0
    for index,row in dfall.iterrows():   #遍历dataframe的每一行
        a = ''
        
        timelist = dfall.loc[index,"审判日期"].split('-')
        dfall.loc[index,"审判日期"] = time2chinese(timelist)
        
        if dfall.loc[index,"案由"] == "申请执行人执行异议之诉":
            dfall.loc[index,"案由"] =  "申请执行人执行异议之与"
            
        domain = os.path.abspath(word_path)
        path = os.path.join(domain,''.join([str(int(row[0])),"-",row[1],".docx"])) #构造文件名
        try:
            document = Document(path)
        except Exception: #防止文件不存在
            print("Can not find {}".format(path))
            count += 1
        if document:
            for paragraph in document.paragraphs[:-1]:
                if paragraph.text.split(): #去除空行
                    paragraph.text = paragraph.text.replace(' ','') #去除空格
                    if paragraph.text[-1].isalpha():    #给没有以标点符号结尾的段落添加句号
                         a += paragraph.text+"。"
                    else:
                         a+= paragraph.text
            if document.paragraphs[-1].text =='在线查看此案例':
                pass
            else:
                a+= paragraph.text

            accuser = re.match(r'.*?原告人?[^\u4e00-\u9fa5]*(.*?)([^\u4e00-\u9fa5]|与|等|被告|诉)', a)  # 找出原告被告
            defendant = re.match(r'.*?被告人?[^\u4e00-\u9fa5]*(.*?)([^\u4e00-\u9fa5]|与|等|的)', a)

            if accuser and defendant:
                accuser = accuser.group(1)
                defendant  = defendant.group(1)
            else:
                accuser,defendant= "本文案中没有原告被告","本文案中没有原告被告"
            dfall.loc[index,"原告"],dfall.loc[index,"被告"]=accuser,defendant  #在dataframe中填充原告被告   #读取在标注之前 所以要先全部读取
            count += 1
            print("\r%.2f%%"%((count/counts)*100), end='')

        document = None

    df = dfall.dropna(axis=0)  # 去掉有空字段的数据

    # df.drop_duplicates(keep='first', inplace=True)  # 去重，只保留第一次出现的样本
    df = df.sample(frac=1.0)  # 全部打乱
    cut_idx = int(round(0.1 * df.shape[0]))
    df_test, df_train = df.iloc[:cut_idx], df.iloc[cut_idx:]
    print (df.shape, df_test.shape, df_train.shape)  # (3184, 12) (318, 12) (2866, 12)

    df.to_csv(os.path.join(save_path, "df.csv"), index=False, encoding='utf_8_sig')
    df_test.to_csv(os.path.join(save_path,"df_test.csv"), index=False, encoding='utf_8_sig')
    df_train.to_csv(os.path.join(save_path,"df_train.csv"), index=False, encoding='utf_8_sig')


def built_ner_data(path_ner,word_path,save_path):
    """建立标注文件"""

    count = 0
    dfall = pd.read_csv(path_ner, encoding='utf-8')
    # dfall.loc[3,"裁判人员"]="钟波,胡振元,周铁金"
    counts = len(dfall)
    entity = ['案号','审理法院','审判日期','裁判人员','案件类型','审判程序','文书类型','案由','原告','被告']
    for index,row in dfall.iterrows():   #遍历dataframe的每一行
        a = ''
        if dfall.loc[index,"案由"] == "申请执行人执行异议之诉":
            dfall.loc[index,"案由"] =  "申请执行人执行异议之与"
            
        domain = os.path.abspath(word_path)
        path = os.path.join(domain,''.join([str(int(row[0])),"-",row[1],".docx"])) #构造文件名
        try:
            document = Document(path)
        except Exception: #防止文件不存在
            print("Can not find {}".format(path))

        for paragraph in document.paragraphs[:-1]:
            if paragraph.text.split(): #去除空行
                paragraph.text = paragraph.text.replace(' ','') #去除空格
                if paragraph.text[-1].isalpha():    #给没有以标点符号结尾的段落添加句号
                     a += paragraph.text+"。"
                else:
                     a+= paragraph.text
        if document.paragraphs[-1].text =='在线查看此案例':
            pass
        else:
            a+= paragraph.text
            
        if len(a)>=500:
            a = a[:400]+a[-100:]
        s=a
        try:
            for j,i in list(enumerate(row[2:])):  #根据dataframe标注数据
                if j == 3:
                    for name in i.split(','): #将裁判人员字段分开
                        # print(name)
                        s = re.sub(name,'[@'+name+'#'+entity[j]+'*]',s)
                else:
                    s=re.sub(i,'[@'+i+'#'+entity[j]+'*]',s)

            str2ner_train_data(s, save_path)

        except Exception:
            print(path)
            continue

        count+=1
        #print(path)
        # print(count)
        print("\r%.2f%%"%((count/counts)*100), end='')
            

         


def str2ner_train_data(s,save_path):
    ner_data = []
    result_1 = re.finditer(r'\[\@', s)
    result_2 = re.finditer(r'\*\]', s)
    begin = []
    end = []
    for each in result_1:
        begin.append(each.start())
    for each in result_2:
        end.append(each.end())
    assert len(begin) == len(end)
    i = 0
    j = 0
    while i < len(s):
        if i not in begin:
            ner_data.append([s[i], 'O'])
            i = i + 1
        else:
            ann = s[i + 2:end[j] - 2]
            # print(ann)
            entity, ner = ann.rsplit('#')
            # print(entity,ner)

            if (len(entity) == 1):
                ner_data.append([entity, 'S-' + ner])
            else:
                if (len(entity) == 2):
                    ner_data.append([entity[0], 'B-' + ner])
                    ner_data.append([entity[1], 'E-' + ner])
                else:
                    ner_data.append([entity[0], 'B-' + ner])
                    for n in range(1, len(entity) - 1):
                        ner_data.append([entity[n], 'I-' + ner])
                    ner_data.append([entity[-1], 'E-' + ner])
 
            i = end[j]
            j = j + 1
 
    f = open(save_path, 'a', encoding='utf-8')
    for each in ner_data:
        f.write(each[0] + ' ' + str(each[1]))
        f.write('\n')
    f.write('\n')
    f.close()


def str2ner_train_data_with_seg(a,s, save_path):
    ner_data = []
    result_1 = re.finditer(r'\[\@', s)
    result_2 = re.finditer(r'\*\]', s)
    begin = []
    end = []
    for each in result_1:
        begin.append(each.start())
    for each in result_2:
        end.append(each.end())
    assert len(begin) == len(end)
    i = 0
    j = 0
    while i < len(s):
        if i not in begin:
            ner_data.append([s[i], 'O'])
            i = i + 1
        else:
            ann = s[i + 2:end[j] - 2]
            # print(ann)
            entity, ner = ann.rsplit('#')
            # print(entity,ner)

            if (len(entity) == 1):
                ner_data.append([entity, 'S-' + ner])
            else:
                if (len(entity) == 2):
                    ner_data.append([entity[0], 'B-' + ner])
                    ner_data.append([entity[1], 'E-' + ner])
                else:
                    ner_data.append([entity[0], 'B-' + ner])
                    for n in range(1, len(entity) - 1):
                        ner_data.append([entity[n], 'I-' + ner])
                    ner_data.append([entity[-1], 'E-' + ner])

            i = end[j]
            j = j + 1
    seg_list = jieba.lcut(a)
    
    i = 0
    for word in seg_list:
        if len(word) == 1:
            ner_data[i].extend('O')
        else:
            if len(word) == 2:
                ner_data[i].extend('B')
                ner_data[i+1].extend('E')
            else:
                ner_data[i].extend('B')
                for n in range(1,len(word) - 1):
                    ner_data[i+n].extend('I')
                ner_data[i+len(word)-1].extend('E')
        i = i + len(word)

    f = open(save_path, 'a', encoding='utf-8')
    for each in ner_data:
        f.write(each[0] + ' ' + str(each[1]) + ' ' + each[2])
        f.write('\n')
    f.write('\n')
    f.close()

def time2chinese(timelist):
    "数字时间转化为汉字"
    date_map = {
        0: '〇',
        1: '一',
        2: '二',
        3: '三',
        4: '四',
        5: '五',
        6: '六',
        7: '七',
        8: '八',
        9: '九'
    }


    def chinese2digits(num, type):
        str_num = str(num)
        result = ''
        if type == 0:
            for i in str_num:
                result = '{}{}'.format(result, date_map.get(int(i)))
        if type == 1:
            result = '{}十{}'.format(date_map.get(int(str_num[0])), date_map.get(int(str_num[1])))
        if type == 2:
            result = '十{}'.format(date_map.get(int(str_num[1])))
        if type == 3:
            result = '十'
        if type == 4:
            result = '二十'
        return result


    year =chinese2digits(int(timelist[0]),0)
    temp = year+"年"
    date_month = int(timelist[1])
    if date_month == 10:
        month = chinese2digits(date_month, 3)
        temp = temp+month+'月'
    if date_month > 10:
        month = chinese2digits(date_month, 2)
        temp = temp+month+'月'
    if date_month < 10:
        month = chinese2digits(date_month, 0)
        temp = temp+month+'月'
    date_day = int(timelist[2])
    if date_day < 10:
        day = chinese2digits(date_day, 0)
        return temp+day+'日'
    if 10 < date_day < 20:
        day = chinese2digits(date_day, 2)
        return temp+day+'日'
    if date_day > 20:
        day = chinese2digits(date_day, 1)
        return temp+day+'日'
    if date_day == 10:
        day = chinese2digits(date_day, 3)
        return temp+day+'日'
    if date_day == 20:
        day = chinese2digits(date_day, 4)
        return temp+day+'日'

if __name__ == '__main__':

    # build_csv("D:/corpus/无讼/excel汇总", 'D:/corpus/无讼/word汇总', "D:/corpus/test/")
    built_ner_data("D:/corpus/test.csv","D:/corpus/无讼/word汇总","D:/corpus/testwithseg.txt")
    built_ner_data("D:/corpus/train.csv","D:/corpus/无讼/word汇总","D:/corpus/trainwithseg.txt")